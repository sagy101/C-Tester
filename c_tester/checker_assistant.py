"""LLM-assisted checker suggestion and audit helpers."""

from __future__ import annotations

import base64
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
import html
import json
import logging
import os
import random
import re
import urllib.error
import urllib.request
import zipfile
from typing import Any, Protocol

import pandas as pd

from .output_contract import ContractConfigError, _extract_field, compile_preset
from .checker_variants import generate_checker_variants
from .semantic_grading import available_checker_templates, checker_config_errors, compare_output_with_config
from .verification import AUDIT_RUBRIC_VERSION


DEFAULT_GEMINI_MODEL = "gemini-flash-latest"
MAX_ASSIGNMENT_IMAGES = 8
MAX_AUDIT_TEXT_CHARS = 100_000
CHECKER_SUGGESTION_EVAL_CASES = {
    "common_supported": [
        {
            "task_shape": "one numeric answer such as factorial, sum, max, count, gcd, or lcm",
            "expected_checker": "last_integer",
            "success_criteria": "Ignore prompts and labels, compare the final numeric answer.",
        },
        {
            "task_shape": "ordered integer sequence such as Fibonacci, primes, sorted array, or printed array values",
            "expected_checker": "integer_list",
            "success_criteria": "Compare the extracted integer sequence, preserving order when the assignment requires order.",
        },
        {
            "task_shape": "divisor-list task, including zero-input wording about no divisors",
            "expected_checker": "divisors",
            "success_criteria": "Validate the divisor set from the numeric input instead of trusting prompt numbers.",
        },
        {
            "task_shape": "reverse integer digits",
            "expected_checker": "reverse_integer",
            "success_criteria": "Compare the reversed integer answer at the configured answer position.",
        },
        {
            "task_shape": "textual yes/no or category answer where punctuation/case are not meaningful",
            "expected_checker": "normalized_text",
            "success_criteria": "Normalize harmless case and punctuation only.",
        },
    ],
    "edge_or_unsupported": [
        {
            "task_shape": "exact menu, drawing, or required sentence where formatting is part of the answer",
            "expected_checker": "exact",
            "success_criteria": "Use exact only when semantic normalization would hide a real mistake.",
        },
        {
            "task_shape": "floating point answer requiring tolerance",
            "expected_checker": "output_contract",
            "success_criteria": "Use labeled_number or floats fields with an approx assertion and a justified tolerance.",
        },
        {
            "task_shape": "multi-column table or compound answer with mixed text and numbers",
            "expected_checker": "output_contract",
            "success_criteria": "Capture each required field and assert it against reference output without relying on layout noise.",
        },
        {
            "task_shape": "assignment text and images include multiple questions",
            "expected_checker": "selected_question_only",
            "success_criteria": "Use only the selected question number and ignore neighboring question context.",
        },
    ],
}
CHECKER_AUDIT_EVAL_CASES = {
    "looks_correct_when": [
        "Perfect score, grade text, output, checker result, and Excel fields agree.",
        "Wrong-input count, comments, and assigned score match checker failures whose expected-vs-actual evidence shows genuine content mistakes.",
        "Compile repair succeeded and the final score reflects the configured repair penalty and repair note.",
        "Structural recursion/loop checks were enforced when configured, and the score/comments reflect any structural penalty.",
        "Submission penalties in final fields match naming, RAR, missing-file, or nested-folder notes.",
    ],
    "flagged_when": [
        "Assigned score conflicts with wrong-input count or checker pass/fail evidence.",
        "Excel comments, timeout fields, compile-error flags, or wrong-input fields contradict the grade text.",
        "Final weighted grade does not match question scores, weights, or submission penalties when those fields are present.",
        "Structural check status, structural penalty, grade text, or final comments contradict each other.",
        "Checker configuration is unsupported or appears mismatched to the assignment intent and would hide mistakes.",
        (
            "The checker rejected student output that is semantically equivalent to the expected output. Compare the "
            "expected-vs-actual text in the grade discrepancies yourself: if the failed cases differ only in phrasing, "
            "wording variants (for example \"is not\" versus \"isn't\"), synonyms, labels, punctuation, spacing, or other "
            "formatting the assignment does not explicitly require, the deduction is a checker defect even when the score "
            "is arithmetically consistent with the checker's pass/fail counts."
        ),
    ],
    "uncertain_when": [
        "Reference context is insufficient to know whether text formatting is semantically important.",
        (
            "Source student output, grade text, or expected/reference output is genuinely missing or ends unexpectedly. "
            "Application compaction explicitly reported in evidence metadata is not student-side truncation."
        ),
        "The case depends on assignment-specific intent not visible in the selected-question context.",
    ],
}
AUDIT_RESPONSE_SCHEMA = {
    "type": "object",
    "properties": {
        "verdict": {"type": "string", "enum": ["looks_correct", "flagged", "uncertain"]},
        "semantic_assessment": {
            "type": "string",
            "enum": ["equivalent", "genuine_error", "unclear"],
        },
        "format_requirement": {
            "type": "string",
            "enum": ["explicit", "not_explicit", "unclear"],
        },
        "checker_behavior": {
            "type": "string",
            "enum": ["correct", "false_reject", "false_accept", "unclear"],
        },
        "risk": {"type": "string", "enum": ["low", "medium", "high"]},
        "reason": {"type": "string"},
        "evidence": {"type": "string"},
    },
    "required": [
        "verdict",
        "semantic_assessment",
        "format_requirement",
        "checker_behavior",
        "risk",
        "reason",
        "evidence",
    ],
    "additionalProperties": False,
}
SUGGEST_CHECKER_RESPONSE_SCHEMA = {
    "type": "object",
    "properties": {
        "status": {"type": "string", "enum": ["supported", "no_supported_checker"]},
        "checker": {"type": "string", "nullable": True},
        "config": {"type": "object"},
        "structural_requirements": {
            "type": "object",
            "properties": {
                "requires_recursion": {"type": "boolean", "nullable": True},
                "entry_functions": {"type": "array", "items": {"type": "string"}, "nullable": True},
                "allow_recursive_helpers": {"type": "boolean", "nullable": True},
                "forbid_loops": {"type": "boolean", "nullable": True},
                "deduction": {"type": "number", "nullable": True},
                "deduction_required": {"type": "boolean", "nullable": True},
                "reason": {"type": "string", "nullable": True},
            },
            "nullable": True,
        },
        "confidence": {"type": "number"},
        "reason": {"type": "string"},
    },
    "required": ["status", "checker", "config", "confidence", "reason"],
    "additionalProperties": False,
}


def get_google_api_key() -> str | None:
    key = os.getenv("GOOGLE_API_KEY")
    if key:
        return key
    if os.name != "nt":
        return None

    try:
        import winreg
    except ImportError:
        return None

    locations = (
        (winreg.HKEY_CURRENT_USER, "Environment"),
        (winreg.HKEY_LOCAL_MACHINE, r"SYSTEM\CurrentControlSet\Control\Session Manager\Environment"),
    )
    for root, subkey in locations:
        try:
            with winreg.OpenKey(root, subkey) as env_key:
                key, _ = winreg.QueryValueEx(env_key, "GOOGLE_API_KEY")
                if key:
                    os.environ["GOOGLE_API_KEY"] = key
                    return key
        except OSError:
            continue
    return None


@dataclass(frozen=True)
class SuggestionResult:
    status: str
    question: str
    checker: str | None
    config: dict
    confidence: float
    reason: str
    structural_requirements: dict | None = None


@dataclass(frozen=True)
class AuditCase:
    student_id: str
    question: str
    score: float
    grade_text: str
    output_text: str
    excel_fields: dict[str, Any]
    final_fields: dict[str, Any]
    reference_output_text: str = ""


@dataclass(frozen=True)
class AuditResult:
    student_id: str
    question: str
    status: str
    verdict: str
    risk: str
    reason: str
    semantic_assessment: str = "unclear"
    format_requirement: str = "unclear"
    checker_behavior: str = "unclear"
    evidence: str = ""
    verification_passes: int = 1
    examples: tuple[dict[str, str], ...] = ()


@dataclass(frozen=True)
class AssignmentImage:
    label: str
    mime_type: str
    data: bytes
    text: str = ""


@dataclass(frozen=True)
class AssignmentContext:
    text: str = ""
    images: tuple[AssignmentImage, ...] = ()
    source_path: str = ""


class LLMProvider(Protocol):
    def complete_json(
        self,
        prompt: str,
        images: list[AssignmentImage] | tuple[AssignmentImage, ...] | None = None,
        response_schema: dict | None = None,
    ) -> dict:
        """Return a JSON object from an LLM prompt."""


class GeminiProvider:
    """Small Gemini REST client using GOOGLE_API_KEY by default."""

    def __init__(self, api_key: str | None = None, model: str | None = None):
        self.api_key = api_key or get_google_api_key()
        self.model = model or os.getenv("GEMINI_MODEL", DEFAULT_GEMINI_MODEL)
        if not self.api_key:
            raise ValueError("GOOGLE_API_KEY is not set")

    def complete_json(
        self,
        prompt: str,
        images: list[AssignmentImage] | tuple[AssignmentImage, ...] | None = None,
        response_schema: dict | None = None,
    ) -> dict:
        url = (
            f"https://generativelanguage.googleapis.com/v1beta/models/{self.model}:"
            f"generateContent?key={self.api_key}"
        )
        parts = [{"text": prompt}]
        for image in images or []:
            parts.append(
                {
                    "inline_data": {
                        "mime_type": image.mime_type,
                        "data": base64.b64encode(image.data).decode("ascii"),
                    }
                }
            )
        generation_config = {"responseMimeType": "application/json"}
        if response_schema:
            generation_config["responseSchema"] = gemini_response_schema(response_schema)
        payload = {
            "contents": [{"parts": parts}],
            "generationConfig": generation_config,
        }
        request = urllib.request.Request(
            url,
            data=json.dumps(payload).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        try:
            with urllib.request.urlopen(request, timeout=60) as response:
                body = json.loads(response.read().decode("utf-8"))
        except urllib.error.HTTPError as exc:
            error_body = exc.read().decode("utf-8", errors="replace")
            raise RuntimeError(f"Gemini request failed: {exc}; {error_body[:1000]}") from exc
        except urllib.error.URLError as exc:
            raise RuntimeError(f"Gemini request failed: {exc}") from exc

        text = body["candidates"][0]["content"]["parts"][0]["text"]
        return parse_json_object(text)


def gemini_response_schema(schema: Any) -> Any:
    """Return the subset of JSON schema accepted by Gemini responseSchema."""
    unsupported_keys = {"additionalProperties", "$schema"}
    if isinstance(schema, dict):
        return {
            key: gemini_response_schema(value)
            for key, value in schema.items()
            if key not in unsupported_keys
        }
    if isinstance(schema, list):
        return [gemini_response_schema(item) for item in schema]
    return schema


def list_gemini_models(api_key: str | None = None) -> list[str]:
    key = api_key or get_google_api_key()
    if not key:
        raise ValueError("GOOGLE_API_KEY is not set")
    url = f"https://generativelanguage.googleapis.com/v1beta/models?key={key}"
    request = urllib.request.Request(url, method="GET")
    try:
        with urllib.request.urlopen(request, timeout=30) as response:
            body = json.loads(response.read().decode("utf-8"))
    except urllib.error.URLError as exc:
        raise RuntimeError(f"Gemini model list request failed: {exc}") from exc
    model_names = []
    for model in body.get("models", []):
        methods = model.get("supportedGenerationMethods", [])
        if "generateContent" not in methods:
            continue
        name = str(model.get("name", "")).removeprefix("models/")
        if name:
            model_names.append(name)
    return sorted(model_names)


def complete_json_with_schema(
    provider: LLMProvider,
    prompt: str,
    images: list[AssignmentImage] | tuple[AssignmentImage, ...] | None = None,
    response_schema: dict | None = None,
) -> dict:
    if response_schema is None:
        return provider.complete_json(prompt, images)
    try:
        return provider.complete_json(prompt, images, response_schema=response_schema)
    except TypeError as exc:
        if "response_schema" not in str(exc):
            raise
        return provider.complete_json(prompt, images)


class FakeLLMProvider:
    """Deterministic provider for tests and offline demos."""

    def complete_json(
        self,
        prompt: str,
        _images: list[AssignmentImage] | tuple[AssignmentImage, ...] | None = None,
        response_schema: dict | None = None,
    ) -> dict:
        del response_schema
        lower_prompt = prompt.lower()
        if '"task": "compile_fix"' in lower_prompt:
            try:
                payload = json.loads(prompt)
                original_code = str(payload.get("original_code", ""))
                current_error = str(payload.get("current_compile_error", ""))
            except json.JSONDecodeError:
                original_code = prompt
                current_error = prompt
            if "too_bad" in original_code.lower() or "ambiguous" in current_error.lower():
                return {
                    "status": "too_bad",
                    "too_bad": True,
                    "fixed_code": "",
                    "compile_issue": "The compile failure is tied to unclear missing logic.",
                    "fix_reason": "code cannot be made to compile without guessing the student's intended logic.",
                    "changes_made": "",
                    "risk_note": "Compile-only repair is unsafe.",
                }
            fixed_code = original_code.replace("/* FAKE_FIX_SEMICOLON */", ";")
            fixed_code = fixed_code.replace("return 0\n}", "return 0;\n}")
            fixed_code = fixed_code.replace("return 0\r\n}", "return 0;\n}")
            return {
                "status": "fixed_candidate",
                "too_bad": False,
                "fixed_code": fixed_code,
                "compile_issue": "A compile-only syntax fix is needed.",
                "fix_reason": "added missing syntax required for compilation.",
                "changes_made": "added missing semicolon",
                "risk_note": "",
            }
        if '"task": "audit_score"' in lower_prompt:
            return {
                "verdict": "uncertain",
                "semantic_assessment": "unclear",
                "format_requirement": "unclear",
                "checker_behavior": "unclear",
                "risk": "medium",
                "reason": "Fake mode cannot verify real grading evidence.",
                "evidence": "No independent semantic reviewer was used.",
            }
        if '"task": "review_score_deduction"' in lower_prompt or '"task": "apply_review_fix"' in lower_prompt:
            return self._fake_review_response(prompt, lower_prompt)
        if "divisor" in lower_prompt:
            return {
                "status": "supported",
                "checker": "divisors",
                "config": {"allow_prompt_numbers": True, "zero_requires_no_divisors_message": True},
                "confidence": 0.9,
                "reason": "The expected outputs describe divisors.",
                "structural_requirements": {},
            }
        if "reverse" in lower_prompt:
            return {
                "status": "supported",
                "checker": "reverse_integer",
                "config": {"answer_position": "last_integer"},
                "confidence": 0.9,
                "reason": "The expected outputs describe reversing an integer.",
                "structural_requirements": {},
            }
        return {
            "status": "no_supported_checker",
            "checker": None,
            "config": {},
            "confidence": 0.0,
            "reason": "No deterministic fake match found.",
            "structural_requirements": {},
        }

    def _fake_review_response(self, prompt: str, lower_prompt: str) -> dict:
        if '"task": "apply_review_fix"' in lower_prompt:
            return self._fake_apply_review_fix_response(prompt)
        return self._fake_score_review_response(prompt)

    @staticmethod
    def _fake_apply_review_fix_response(prompt: str) -> dict:
        try:
            payload = json.loads(prompt)
            current_code = str(payload.get("current_code", ""))
        except json.JSONDecodeError:
            current_code = prompt
        return {
            "fixed_code": "/* REVIEWER FIX: fake provider left code unchanged for offline preview. */\n" + current_code,
            "explanation": "Fake/offline mode demonstrates the fix workflow without changing program logic.",
            "changes_made": ["Added a visible reviewer-fix comment only."],
            "tests_to_run": ["Run all grading inputs and inspect any remaining mismatches."],
            "risk_note": "Fake/offline mode does not produce real code fixes.",
        }

    @staticmethod
    def _fake_score_review_response(prompt: str) -> dict:
        try:
            payload = json.loads(prompt)
            failed_cases = payload.get("failed_cases", [])
        except json.JSONDecodeError:
            failed_cases = []
        failed_inputs = [str(case.get("input_value", "")) for case in failed_cases if isinstance(case, dict)]
        examples = [
            {
                "input": str(case.get("input_value", "")),
                "expected_output": str(case.get("expected_output", "")),
                "actual_output": str(case.get("actual_output", "")),
                "why_it_failed": str(case.get("reason", "")) or "The actual output did not match the expected output.",
            }
            for case in failed_cases[:3]
            if isinstance(case, dict)
        ]
        return {
            "summary": "The fake reviewer grouped the supplied failures as one deterministic grading issue.",
            "deduction_is_plausible": True,
            "deduction_caused_by": "student_code",
            "semantic_assessment": "genuine_error",
            "format_requirement": "unclear",
            "format_requirement_evidence": "",
            "root_causes": [
                {
                    "issue": "The output differs from the expected format or value for the listed inputs.",
                    "failed_inputs": failed_inputs,
                    "deduction_impact": f"{len(failed_inputs)} failed input(s) contributed to the shown deduction.",
                    "examples": examples,
                }
            ],
            "inline_comments": [
                {
                    "line": 1,
                    "comment": "Start by checking the branch or formatting responsible for the first failed case.",
                }
            ],
            "fix_to_full_score": "Make the output match the expected behavior for every failed input.",
            "risk_note": "",
        }


def parse_json_object(text: str) -> dict:
    text = strip_json_fence(text.strip())
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError as exc:
        candidate = extract_first_json_object_text(text)
        if not candidate:
            raise exc
        parsed = json.loads(candidate)
    if not isinstance(parsed, dict):
        raise ValueError("LLM response must be a JSON object")
    return parsed


def strip_json_fence(text: str) -> str:
    if not text.startswith("```"):
        return text
    lines = text.splitlines()
    if lines and lines[0].strip().startswith("```"):
        lines = lines[1:]
    if lines and lines[-1].strip() == "```":
        lines = lines[:-1]
    return "\n".join(lines).strip()


def extract_first_json_object_text(text: str) -> str:
    start = text.find("{")
    if start == -1:
        return ""
    depth = 0
    in_string = False
    escaped = False
    for index in range(start, len(text)):
        char = text[index]
        if in_string or char == '"':
            in_string, escaped, consumed = advance_json_string_state(char, in_string, escaped)
            if consumed:
                continue
        if char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return text[start : index + 1]
    return ""


def advance_json_string_state(char: str, in_string: bool, escaped: bool) -> tuple[bool, bool, bool]:
    if not in_string:
        return True, False, True
    if escaped:
        return True, False, True
    if char == "\\":
        return True, True, True
    if char == '"':
        return False, False, True
    return True, False, True


def parse_assignment_file(path: str | None) -> str:
    return parse_assignment_context(path).text


def parse_assignment_context(path: str | None) -> AssignmentContext:
    if not path:
        return AssignmentContext()
    if not os.path.exists(path):
        raise FileNotFoundError(path)
    ext = os.path.splitext(path)[1].lower()
    if ext in {".txt", ".md", ".csv"}:
        with open(path, "r", encoding="utf-8", errors="ignore") as text_file:
            return AssignmentContext(text=text_file.read(), source_path=path)
    if ext == ".docx":
        return parse_docx_context(path)
    if ext == ".pdf":
        return parse_pdf_context(path)
    with open(path, "r", encoding="utf-8", errors="ignore") as text_file:
        return AssignmentContext(text=text_file.read(), source_path=path)


def parse_docx_text(path: str) -> str:
    return parse_docx_context(path).text


def parse_docx_context(path: str) -> AssignmentContext:
    text = parse_docx_text_with_package(path)
    images = extract_docx_images(path)
    return AssignmentContext(text=text, images=tuple(images), source_path=path)


def parse_docx_text_with_package(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        return parse_docx_text_from_xml(path)

    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def parse_docx_text_from_xml(path: str) -> str:
    with zipfile.ZipFile(path) as docx:
        xml = docx.read("word/document.xml").decode("utf-8", errors="ignore")
    text = re.sub(r"<[^>]+>", " ", xml)
    return html.unescape(" ".join(text.split()))


def extract_docx_images(path: str) -> list[AssignmentImage]:
    images = []
    with zipfile.ZipFile(path) as docx:
        for name in docx.namelist():
            if not name.startswith("word/media/"):
                continue
            mime_type = image_mime_type(name)
            if not mime_type:
                continue
            images.append(
                AssignmentImage(
                    label=f"DOCX embedded image {len(images) + 1}: {os.path.basename(name)}",
                    mime_type=mime_type,
                    data=docx.read(name),
                )
            )
            if len(images) >= MAX_ASSIGNMENT_IMAGES:
                break
    return images


def parse_pdf_text_best_effort(path: str) -> str:
    return parse_pdf_context(path).text


def parse_pdf_context(path: str) -> AssignmentContext:
    try:
        return parse_pdf_context_with_pymupdf(path)
    except ImportError:
        text = parse_pdf_text_with_pypdf(path)
        return AssignmentContext(text=text, source_path=path)


def parse_pdf_context_with_pymupdf(path: str) -> AssignmentContext:
    try:
        import fitz
    except ImportError as exc:
        raise ImportError("PDF image extraction requires PyMuPDF. Run: python -m pip install pymupdf") from exc

    document = fitz.open(path)
    text_parts = []
    images = []
    for page_index, page in enumerate(document, start=1):
        page_text = page.get_text("text").strip()
        if page_text:
            text_parts.append(page_text)
        if len(images) < MAX_ASSIGNMENT_IMAGES:
            pixmap = page.get_pixmap(matrix=fitz.Matrix(1.25, 1.25), alpha=False)
            images.append(
                AssignmentImage(
                    label=f"PDF page {page_index}",
                    mime_type="image/png",
                    data=pixmap.tobytes("png"),
                    text=page_text,
                )
            )
    text = "\n".join(text_parts).strip()
    if not text:
        text = parse_pdf_text_with_pypdf(path)
    return AssignmentContext(text=text, images=tuple(images), source_path=path)


def parse_pdf_text_with_pypdf(path: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF assignment parsing requires pypdf. Run: python -m pip install pypdf") from exc

    logging.getLogger("pypdf").setLevel(logging.ERROR)
    reader = PdfReader(path, strict=False)
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(page_text.strip() for page_text in pages if page_text.strip())
    if not text:
        raise ValueError("Could not extract text from PDF assignment file.")
    return text


def image_mime_type(path: str) -> str | None:
    ext = os.path.splitext(path)[1].lower()
    return {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".webp": "image/webp",
    }.get(ext)


def assignment_context_for_question(context: AssignmentContext | str, question: str) -> AssignmentContext:
    if isinstance(context, str):
        context = AssignmentContext(text=context)
    question_text = extract_question_assignment_text(context.text, question)
    images = tuple(image for image in context.images if assignment_image_matches_question(image, question))
    if context.images and not images:
        images = context.images[:MAX_ASSIGNMENT_IMAGES]
    return AssignmentContext(
        text=question_text,
        images=images[:MAX_ASSIGNMENT_IMAGES],
        source_path=context.source_path,
    )


def extract_question_assignment_text(text: str, question: str) -> str:
    if not text:
        return ""
    question_number = question_number_from_name(question)
    if question_number is None:
        return text[:12000]

    matches = list(question_heading_matches(text))
    preamble = text[: matches[0][1]].strip() if matches else ""
    for index, match in enumerate(matches):
        if match[0] != question_number:
            continue
        start = match[1]
        end = matches[index + 1][1] if index + 1 < len(matches) else len(text)
        question_text = text[start:end].strip()
        if preamble:
            return f"Global assignment instructions:\n{preamble}\n\nSelected question:\n{question_text}"[:12000]
        return question_text[:12000]
    return text[:12000]


def assignment_image_matches_question(image: AssignmentImage, question: str) -> bool:
    question_number = question_number_from_name(question)
    if question_number is None or not image.text:
        return not image.text
    return any(match[0] == question_number for match in question_heading_matches(image.text))


def question_number_from_name(question: str) -> int | None:
    match = re.search(r"\d+", question or "")
    return int(match.group(0)) if match else None


def question_heading_matches(text: str):
    pattern = re.compile(r"(?:^|\n)\s*(?:question\s*(\d+)\b|q\s*(\d+)\b|שאלה\s*(\d+))", re.IGNORECASE)
    for match in pattern.finditer(text):
        number_text = match.group(1) or match.group(2) or match.group(3)
        if number_text:
            yield int(number_text), match.start()


def build_suggestion_prompt(
    question: str,
    original_code: str,
    inputs: list[str],
    expected_outputs: list[tuple[str, str]],
    assignment_text: str = "",
    assignment_images: list[AssignmentImage] | tuple[AssignmentImage, ...] | None = None,
) -> str:
    question_number = question_number_from_name(question)
    payload = {
        "task": "suggest_checker",
        "role": "You are configuring a deterministic C homework output checker. You do not grade students and you do not write code.",
        "question": question,
        "target_question_number": question_number,
        "available_checkers": available_checker_templates(),
        "output_contract_language": {
            "shape": {
                "checker": "output_contract",
                "config": {
                    "contract": {
                        "version": 1,
                        "description": "assignment-neutral description",
                        "fields": "list of field objects",
                        "checks": "list of assertion objects",
                    }
                },
            },
            "field_keys": [
                "id", "source", "extract", "normalize", "select", "label", "labels", "number_type",
                "anchor", "anchors", "occurrence", "count", "window", "true_aliases", "false_aliases",
            ],
            "sources": ["stdin", "reference", "actual"],
            "extractors": ["text", "integers", "floats", "labeled_number", "point", "points", "boolean"],
            "extractor_requirements": {
                "labeled_number": "requires label/labels, or anchor/anchors to capture the first following number",
                "point": "optional anchor; occurrence defaults to 0; returns one coordinate pair",
                "points": "optional anchor; count is required when more than two points matter",
                "boolean": (
                    "requires non-empty true_aliases and false_aliases; anchor is strongly recommended. The earliest "
                    "alias occurrence wins and longer aliases win position ties, so cover negated phrasing variants in "
                    "false_aliases (for example both \"isn't\" and \"is not\") rather than relying on a bare positive "
                    "alias. An alias directly followed by a negation (\"not\"/\"n't\") is automatically inverted."
                ),
                "integers_or_floats": "select is all, last, {index: n}, or {slice: [start, count]}",
            },
            "normalizers": ["collapse_whitespace", "lowercase", "strip_punctuation", "normalize_apostrophe"],
            "selectors": ["all", "last", {"index": 0}, {"slice": [0, 2]}],
            "check_shape": {
                "id": "stable_generic_id",
                "op": "equal | approx | sequence_equal | tail_equal | exchanged",
                "left": {"field": "actual_field_id"},
                "right": {"field": "reference_or_stdin_field_id"},
                "tolerance": "non-negative number when needed",
                "ordered": "boolean for sequences",
                "message": "short failure reason",
            },
            "safety_rules": [
                "Do not emit regex, source code, formulas, expressions, or unknown keys.",
                "Use reference fields for assignment-specific calculations.",
                "Use stdin fields and exchanged assertions for generic state-transition invariants.",
                "Use anchors only as short literal phrases present in both intended output variants.",
                "Every semantically required output fact needs its own field and check.",
                "For programs printing before/after values, check both values with point/points fields and exchanged assertions; address text alone is not enough.",
                "Cover every function behavior exercised by the reference main program, not only final calculations.",
                "Numeric tolerance cannot exceed 0.011. Use 0.011 for two-decimal output and a smaller value for higher precision.",
                (
                    "Students phrase equivalent facts differently. Configure aliases and anchors so semantically "
                    "equivalent wording passes: include common contractions and their expansions in boolean aliases, "
                    "and never anchor on decorative text a correct program could reasonably omit or reword."
                ),
            ],
        },
        "eval_cases": CHECKER_SUGGESTION_EVAL_CASES,
        "original_solution_c": original_code,
        "inputs": inputs,
        "expected_outputs": [{"input": value, "output": output} for value, output in expected_outputs],
        "assignment_text_for_selected_question_optional": assignment_text,
        "assignment_images_for_selected_question_optional": [image.label for image in assignment_images or []],
        "instructions": (
            f"Focus only on the selected question {question} (question number {question_number}). If the assignment text "
            "or images contain multiple questions, ignore the other question sections. Infer the intended answer and output "
            "format from the selected-question assignment context when present, the reference C code, the test inputs, and "
            "the reference outputs. Select a generic preset when it fully captures the answer; otherwise emit an "
            "output_contract using only the supplied declarative language. Prefer the simplest checker that validates every "
            "semantic requirement while ignoring harmless prompts/labels/spacing. Do not invent code, regular expressions, "
            "formulas, checker names, keys, extractors, or operators. "
            "Each input string is the full stdin sent to the program, not necessarily the mathematical argument. "
            "Some assignments use routing/menu fields before the actual question argument. "
            "For those multi-integer/menu inputs, do not choose checkers that derive the answer directly from the raw stdin "
            "unless that checker can be configured with input_integer_index to point at the actual task argument. "
            "If no input-derived checker safely matches the task, compare the reference output answer, usually with "
            "last_integer for one numeric Result value. "
            "If no available checker can safely validate the task, return status no_supported_checker. When uncertain, "
            "lower confidence and explain what manual review is needed. If the selected question requires recursion, "
            "forbids loops, or gives an explicit structural penalty, include structural_requirements. Derive deduction "
            "only from explicit assignment text such as 'non-recursive gets 0'. If the assignment requires recursion "
            "but does not state the deduction, set deduction to null and deduction_required to true so the grader can "
            "fill the mandatory value manually."
        ),
        "decision_guidance": {
            "exact": "Use only when exact output format is required.",
            "normalized_text": "Use for mostly textual answers where case/punctuation are irrelevant.",
            "last_integer": (
                "Use for tasks with one numeric answer at the end, such as factorial, sum, max, count, gcd, lcm, "
                "binary-as-digits, digit reductions, or a series value. Prefer this for menu/routing-style stdin "
                "when the final answer is printed as Result = <number>."
            ),
            "integer_list": "Use for sequences/lists like Fibonacci, primes, sorted arrays, or printed array values.",
            "divisors": (
                "Use only for divisor-list tasks when the raw stdin value is the number whose divisors are expected. "
                "If stdin includes routing/menu fields, set input_integer_index to the argument integer. "
                "Do not use it for unrelated numeric transformations."
            ),
            "reverse_integer": (
                "Legacy alias for final integer comparison; prefer last_integer for new configurations. "
                "Do not use it for digit-sum/digit-reduction tasks. Historical input_integer_index settings for "
                "routing/menu fields are accepted for compatibility but are not needed by the reference-output contract."
            ),
            "output_contract": (
                "Use for labeled floats, tolerances, mixed fields, sections, aliases, and before/after traces. "
                "Compare assignment-specific calculations to reference fields and express only generic invariants."
            ),
            "no_supported_checker": "Use only when the safe declarative language cannot represent an essential requirement.",
        },
        "response_schema": {
            "status": "supported | no_supported_checker",
            "checker": "checker name or null",
            "config": "object",
            "structural_requirements": {
                "requires_recursion": "boolean, true only when assignment requires recursion",
                "entry_functions": "array of function names to inspect, e.g. ['q_2']",
                "allow_recursive_helpers": "boolean, true when helper recursion is acceptable",
                "forbid_loops": "boolean, true when loops are disallowed",
                "deduction": "number of points to deduct, or null when mandatory manual input is needed",
                "deduction_required": "boolean, true when deduction could not be derived from assignment text",
                "reason": "short explanation grounded in the assignment text",
            },
            "confidence": "0.0-1.0",
            "reason": "short explanation of why this checker is safe and what output it validates",
        },
    }
    return json.dumps(payload, ensure_ascii=False, indent=2)


def suggest_checker(
    question: str,
    original_code: str,
    inputs: list[str],
    expected_outputs: list[tuple[str, str]],
    provider: LLMProvider,
    assignment_text: str = "",
    assignment_images: list[AssignmentImage] | tuple[AssignmentImage, ...] | None = None,
) -> SuggestionResult:
    prompt = build_suggestion_prompt(question, original_code, inputs, expected_outputs, assignment_text, assignment_images)
    # The declarative contract contains heterogeneous field/check objects that
    # Gemini's restricted responseSchema dialect cannot faithfully represent.
    # JSON mode is still enforced by the provider; the result is then validated
    # by the deterministic contract validator before testing or saving.
    response = _complete_checker_json_with_retry(provider, prompt, assignment_images)
    suggestion = _suggestion_from_response(question, response, assignment_text)
    if suggestion.status == "supported" and suggestion.checker:
        errors = checker_config_errors({"checker": suggestion.checker, "config": suggestion.config})
        if errors:
            retry_prompt = (
                f"{prompt}\n\n"
                "Your previous candidate was rejected by deterministic schema validation:\n"
                + "\n".join(f"- {error}" for error in errors)
                + "\nReturn a corrected complete JSON response. Do not omit required fields or checks."
            )
            response = _complete_checker_json_with_retry(provider, retry_prompt, assignment_images)
            suggestion = _suggestion_from_response(question, response, assignment_text)
    return suggestion


def _suggestion_from_response(question: str, response: dict, assignment_text: str) -> SuggestionResult:
    structural_requirements = merge_structural_requirements(
        question,
        response.get("structural_requirements") if isinstance(response.get("structural_requirements"), dict) else {},
        assignment_text,
    )
    return SuggestionResult(
        status=str(response.get("status", "no_supported_checker")),
        question=question,
        checker=response.get("checker"),
        config=response.get("config") if isinstance(response.get("config"), dict) else {},
        confidence=float(response.get("confidence", 0.0) or 0.0),
        reason=str(response.get("reason", "")),
        structural_requirements=structural_requirements,
    )


def refine_checker(
    question: str,
    original_code: str,
    inputs: list[str],
    expected_outputs: list[tuple[str, str]],
    current_config: dict,
    audit_results: list[AuditResult],
    provider: LLMProvider,
    assignment_text: str = "",
    assignment_images: list[AssignmentImage] | tuple[AssignmentImage, ...] | None = None,
    review_feedback: list[dict] | None = None,
) -> SuggestionResult:
    del original_code, inputs
    base_payload = {
        "task": "refine_checker",
        "question": question,
        "target_question_number": question_number_from_name(question),
        "current_checker_config": current_config,
        "assignment_text_for_selected_question_optional": str(assignment_text)[:8000],
        "reference_examples": [
            {"input": input_value, "output": output}
            for input_value, output in expected_outputs[:3]
        ],
        "contract_language": {
            "field_keys": [
                "id", "source", "extract", "normalize", "select", "label", "labels",
                "number_type", "anchor", "anchors", "occurrence", "count", "window",
                "true_aliases", "false_aliases", "allow_empty",
            ],
            "extractors": ["text", "integers", "floats", "labeled_number", "point", "points", "boolean"],
            "operators": ["equal", "approx", "sequence_equal", "tail_equal", "sequence_or_text_tokens", "exchanged"],
            "limits": "Keep tolerance <= 0.011; labels/anchors/aliases are bounded literal strings, never regex.",
        },
        "response_shape": {
            "status": "supported | no_supported_checker",
            "checker": "checker name or null",
            "config": "complete config object",
            "structural_requirements": "preserve current requirements",
            "confidence": "0.0-1.0",
            "reason": "short explanation",
        },
    }
    base_payload["audit_feedback"] = [
        {
            "status": result.status,
            "verdict": result.verdict,
            "risk": result.risk,
            "reason": result.reason,
            "semantic_assessment": result.semantic_assessment,
            "format_requirement": result.format_requirement,
            "checker_behavior": result.checker_behavior,
            "evidence": result.evidence,
            "examples": list(result.examples),
        }
        for result in audit_results
        if result.status in {"flagged", "uncertain"}
    ]
    review_items = [
        item
        for item in (review_feedback or [])
        if corroborated_review_feedback_item(item)
    ]
    base_payload["review_feedback"] = [
        {
            "cause": "checker_or_app",
            "summary": str(item.get("summary", "")),
            "risk_note": str(item.get("risk_note", "")),
            "student_label": str(item.get("anonymized_label") or item.get("student_id") or "student"),
            "semantic_assessment": str(item.get("semantic_assessment", "unclear")),
            "format_requirement": str(item.get("format_requirement", "unclear")),
            "format_requirement_evidence": str(item.get("format_requirement_evidence", "")),
            "root_causes": item.get("root_causes", []) if isinstance(item.get("root_causes", []), list) else [],
        }
        for item in review_items
    ]
    base_payload["instructions"] = (
        "Return the complete checker JSON, not a patch. Improve the current checker only where audit_feedback or "
        "review_feedback identifies a checker false "
        "rejection or false acceptance. Review findings with cause=checker_or_app are high-confidence evidence that "
        "the checker penalized semantically equivalent student output; fix aliases, anchors, labels, and operators so "
        "those equivalent phrasings pass while genuine content mistakes still fail. Preserve every semantic "
        "requirement already checked. Return the unchanged configuration when feedback is uncertainty, missing "
        "evidence, or a student-specific grading issue rather than a checker defect. The candidate will be rejected "
        "unless it passes positive invariance, semantic mutation, cumulative no-regression, and audited-student gates. "
        "Use labels/anchors arrays only for bounded confirmed alternatives; never make extraction unrestricted."
    )
    response = _complete_checker_json_with_retry(
        provider,
        json.dumps(base_payload, ensure_ascii=False, indent=2),
        assignment_images,
    )
    structural_requirements = merge_structural_requirements(
        question,
        response.get("structural_requirements") if isinstance(response.get("structural_requirements"), dict) else {},
        assignment_text,
    )
    return SuggestionResult(
        status=str(response.get("status", "no_supported_checker")),
        question=question,
        checker=response.get("checker"),
        config=response.get("config") if isinstance(response.get("config"), dict) else {},
        confidence=float(response.get("confidence", 0.0) or 0.0),
        reason=str(response.get("reason", "")),
        structural_requirements=structural_requirements,
    )


def _complete_checker_json_with_retry(
    provider: LLMProvider,
    prompt: str,
    images: list[AssignmentImage] | tuple[AssignmentImage, ...] | None = None,
) -> dict:
    try:
        return complete_json_with_schema(provider, prompt, images)
    except Exception as first_error:
        retry_prompt = (
            f"{prompt}\n\n"
            "Your previous response failed transport or JSON parsing. Return one complete valid JSON object only, "
            "with no markdown fence, comments, or trailing text."
        )
        try:
            return complete_json_with_schema(provider, retry_prompt, images)
        except Exception as second_error:
            raise RuntimeError(f"{first_error}; checker JSON retry failed: {second_error}") from second_error


def corroborated_review_feedback_item(item: Any) -> bool:
    if not isinstance(item, dict):
        return False
    cause = str(item.get("cause", item.get("deduction_caused_by", ""))).strip()
    semantic = str(item.get("semantic_assessment", "unclear")).strip()
    format_requirement = str(item.get("format_requirement", "unclear")).strip()
    root_causes = item.get("root_causes", [])
    has_example = any(
        isinstance(root, dict)
        and any(
            isinstance(example, dict)
            and bool(example.get("expected_output"))
            and bool(example.get("actual_output"))
            for example in root.get("examples", [])
        )
        for root in root_causes
        if isinstance(root_causes, list)
    )
    return (
        cause == "checker_or_app"
        and semantic == "equivalent"
        and format_requirement != "explicit"
        and has_example
    )


def review_feedback_test_rows(checker_config: dict, items: list[dict] | None) -> list[dict]:
    """Create accept rows and paired semantic corruptions from corroborated reviews."""
    rows: list[dict] = []
    for item in items or []:
        if not corroborated_review_feedback_item(item):
            continue
        for root in item.get("root_causes", []):
            if not isinstance(root, dict):
                continue
            for example in root.get("examples", []):
                if not isinstance(example, dict):
                    continue
                input_value = str(example.get("input", ""))
                expected = str(example.get("expected_output", ""))
                actual = str(example.get("actual_output", ""))
                if not expected or not actual:
                    continue
                rows.append(
                    {
                        "variant": "accept_adjudicated_student_variant",
                        "input": input_value,
                        "expected_output": expected,
                        "actual_output": actual,
                        "expected_pass": True,
                    }
                )
                for name, corrupted, expected_pass in generate_checker_variants(checker_config, actual):
                    if expected_pass:
                        continue
                    rows.append(
                        {
                            "variant": f"paired_{name}",
                            "input": input_value,
                            "expected_output": expected,
                            "actual_output": corrupted,
                            "expected_pass": False,
                        }
                    )
                if len(rows) >= 40:
                    return rows
    return rows


def merge_structural_requirements(question: str, existing: dict | None, assignment_text: str = "") -> dict:
    merged = dict(existing or {})
    inferred = infer_structural_requirements(question, assignment_text)
    if not inferred:
        normalize_structural_requirement_flags(merged)
        return merged

    if not (merged.get("requires_recursion") or merged.get("forbid_loops")):
        normalize_structural_requirement_flags(inferred)
        return inferred

    for key, value in inferred.items():
        if merged.get(key) in (None, "", [], {}):
            merged[key] = value
    normalize_structural_requirement_flags(merged)
    return merged


def normalize_structural_requirement_flags(requirements: dict):
    deduction = requirements.get("deduction", requirements.get("penalty"))
    try:
        float(deduction)
    except (TypeError, ValueError):
        return
    requirements["deduction_required"] = False


def infer_structural_requirements(question: str, assignment_text: str = "") -> dict:
    normalized = " ".join(str(assignment_text or "").lower().split())
    has_recursion_requirement = any(
        phrase in normalized
        for phrase in [
            "recursion",
            "recursive",
            "non-recursive",
            "רקורס",
        ]
    )
    if not has_recursion_requirement:
        return {}

    zero_deduction = any(
        phrase in normalized
        for phrase in [
            "will receive credit",
            "will not receive credit",
            "receive 0",
            "gets 0",
            "get 0",
            "יקבל 0",
            "0 !",
            "0!",
        ]
    )
    question_number = question_number_from_name(question)
    requirements = {
        "requires_recursion": True,
        "entry_functions": [f"q_{question_number}"] if question_number is not None else [],
        "allow_recursive_helpers": True,
        "forbid_loops": "loop" in normalized or "for/while" in normalized or "while" in normalized or "for" in normalized,
        "deduction_required": not zero_deduction,
        "reason": "Assignment text requires recursive implementation.",
    }
    if zero_deduction:
        requirements["deduction"] = 100
        requirements["deduction_required"] = False
        requirements["reason"] = "Assignment states non-recursive solutions receive 0."
    else:
        requirements["deduction"] = None
    return requirements


def run_checker_tests(checker_config: dict, inputs_and_expected: list[tuple[str, str]], max_cases: int = 8) -> list[dict]:
    rows = []
    checker_name = (checker_config or {}).get("checker", "exact")
    prompted_expected = checker_name not in {"exact", "normalized_text"}
    for input_value, expected_output in inputs_and_expected[:max_cases]:
        variants = [
            ("exact", expected_output, True),
            ("prompted", f"Input: {input_value}\nOutput: {expected_output}", prompted_expected),
            ("wrong", "__synthetic_output_without_an_answer__", False),
        ]
        variants.extend(generate_checker_variants(checker_config, expected_output))
        for variant_name, actual_output, expected_pass in variants:
            comparison = compare_output_with_config(checker_config, input_value, expected_output, actual_output)
            test_passed = comparison.passed == expected_pass
            rows.append(
                {
                    "input": input_value,
                    "variant": variant_name,
                    "expected_output": expected_output,
                    "actual_output": actual_output,
                    "passed": comparison.passed,
                    "test_passed": test_passed,
                    "comparison_passed": comparison.passed,
                    "expected_pass": expected_pass,
                    "reason": comparison.reason if test_passed else f"Expected {'accept' if expected_pass else 'reject'}: {comparison.reason}",
                    "expected_canonical": comparison.expected_canonical,
                    "actual_canonical": comparison.actual_canonical,
                }
            )
    return rows


def _contract_mutation_variants(checker_config: dict, expected_output: str) -> list[tuple[str, str, bool]]:
    if (checker_config or {}).get("checker") != "output_contract":
        return []
    try:
        contract = compile_preset(
            (checker_config or {}).get("checker", "exact"),
            (checker_config or {}).get("config", {}),
        )
    except ContractConfigError:
        return []

    variants: list[tuple[str, str, bool]] = []
    seen_outputs = set()
    actual_field_ids = {
        check.get(side, {}).get("field")
        for check in contract.get("checks", [])
        for side in ("left", "right")
    }
    for field in contract.get("fields", []):
        if field.get("source") != "actual" or field.get("id") not in actual_field_ids:
            continue
        mutated = _mutate_output_for_field(expected_output, field)
        if (
            mutated
            and mutated != expected_output
            and mutated not in seen_outputs
            and _mutation_changes_extracted_value(field, expected_output, mutated)
        ):
            variants.append((f"reject_{field['id']}", mutated, False))
            seen_outputs.add(mutated)
        if len(variants) >= 6:
            break
    if len(expected_output) > 20:
        truncated = expected_output[: len(expected_output) // 2]
        if truncated not in seen_outputs:
            variants.append(("reject_truncated", truncated, False))
    # Semantically equivalent rewording must keep passing: expanding
    # contractions ("isn't" -> "is not") preserves meaning, so a contract that
    # rejects it would deduct points for phrasing instead of content.
    reworded = _expand_negative_contractions(expected_output)
    if reworded != expected_output and reworded not in seen_outputs:
        variants.append(("accept_equivalent_negation", reworded, True))
    return variants


def _expand_negative_contractions(text: str) -> str:
    def replace(match: re.Match) -> str:
        word = match.group(0)
        # Irregular stems ("can't" -> "ca") would produce gibberish; keep them.
        if word[:-3].lower() in {"ca", "wo", "sha", "ai"}:
            return word
        return f"{word[:-3]} not"

    return re.sub(r"[A-Za-z]+n't\b", replace, text, flags=re.IGNORECASE)


def _mutation_changes_extracted_value(field: dict, original: str, mutated: str) -> bool:
    try:
        return _extract_field(field, original) != _extract_field(field, mutated)
    except (TypeError, ValueError):
        return False


def _mutate_output_for_field(output: str, field: dict) -> str | None:
    extractor = field.get("extract")
    anchor = field.get("anchor")
    start = 0
    if anchor:
        start = output.lower().find(str(anchor).lower())
        if start < 0:
            return None
        start += len(anchor)
    window_end = min(len(output), start + int(field.get("window", 4096)))
    scoped = output[start:window_end]

    if extractor == "labeled_number":
        label = field.get("label")
        match = re.search(
            rf"({re.escape(str(label))}\s*[:=]?\s*)([+-]?(?:\d+(?:\.\d*)?|\.\d+))",
            scoped,
            re.IGNORECASE,
        )
        if not match:
            return None
        old_value = match.group(2)
        replacement = f"{float(old_value) + 1:g}"
        return output[: start + match.start(2)] + replacement + output[start + match.end(2):]

    if extractor == "boolean":
        pairs = [
            (true_alias, false_alias)
            for true_alias in field.get("true_aliases", [])
            for false_alias in field.get("false_aliases", [])[:1]
        ] + [
            (false_alias, true_alias)
            for false_alias in field.get("false_aliases", [])
            for true_alias in field.get("true_aliases", [])[:1]
        ]
        for old, new in pairs:
            position = scoped.lower().find(old.lower())
            if position >= 0:
                absolute = start + position
                return output[:absolute] + new + output[absolute + len(old):]
        return None

    if extractor in {"point", "points", "integers", "floats"}:
        match = re.search(r"[+-]?(?:\d+(?:\.\d*)?|\.\d+)", scoped)
        if not match:
            return None
        replacement = f"{float(match.group(0)) + 1:g}"
        return output[: start + match.start()] + replacement + output[start + match.end():]

    if extractor == "text":
        return output + "\n__unexpected_semantic_value__"
    return None


def select_audit_cases(
    questions: list[str],
    max_cases: int = 15,
    seed: int | None = None,
    exclude_student_ids: set[str] | None = None,
) -> list[AuditCase]:
    cases_by_bucket = _empty_audit_buckets()
    cases_by_question = {question: _empty_audit_buckets() for question in questions}
    final_df = _read_excel_if_exists("final_grades.xlsx")
    final_by_id = _rows_by_id(final_df)
    excluded = {str(student_id) for student_id in (exclude_student_ids or set())}

    for question in questions:
        grade_excel = os.path.join(question, f"{question}_grades_to_upload.xlsx")
        question_df = _read_excel_if_exists(grade_excel)
        for _, row in question_df.iterrows():
            case = _make_audit_case(question, row, final_by_id)
            if case.student_id in excluded:
                continue
            _add_case_to_buckets(cases_by_bucket, case)
            _add_case_to_buckets(cases_by_question[question], case)

    if seed is not None:
        rng = random.Random(seed)
        for buckets in [cases_by_bucket, *cases_by_question.values()]:
            for bucket_cases in buckets.values():
                rng.shuffle(bucket_cases)

    selected = []
    if max_cases >= len(questions):
        for question in questions:
            question_pick = _take_bucketed_cases(cases_by_question[question], 1)
            if question_pick and len(selected) < max_cases:
                selected.append(question_pick[0])

    return _take_bucketed_cases(cases_by_bucket, max_cases, selected)


def _empty_audit_buckets() -> dict[str, list[AuditCase]]:
    return {
        "extraction_failure": [],
        "semantic_mismatch": [],
        "perfect": [],
        "high": [],
        "medium": [],
        "low": [],
        "zero": [],
        "penalty": [],
        "compile_timeout": [],
    }


def _make_audit_case(question: str, row: pd.Series, final_by_id: dict[str, dict]) -> AuditCase:
    student_id = str(row.get("ID_number", ""))
    score = float(row.get("Grade", 0) or 0)
    return AuditCase(
        student_id=student_id,
        question=question,
        score=score,
        grade_text=_read_optional_text(os.path.join(question, "grade", f"{student_id}.txt")),
        output_text=_read_optional_text(os.path.join(question, "output", f"{student_id}.txt")),
        excel_fields=row.to_dict(),
        final_fields=final_by_id.get(student_id, {}),
        reference_output_text=_read_optional_text(os.path.join(question, "original_sol_output.txt")),
    )


def _add_case_to_buckets(cases_by_bucket: dict[str, list[AuditCase]], case: AuditCase):
    for bucket in _case_buckets(case):
        cases_by_bucket[bucket].append(case)


def _take_bucketed_cases(
    cases_by_bucket: dict[str, list[AuditCase]],
    max_cases: int,
    selected: list[AuditCase] | None = None,
) -> list[AuditCase]:
    selected = list(selected or [])
    while len(selected) < max_cases and any(cases_by_bucket.values()):
        for bucket in cases_by_bucket:
            if cases_by_bucket[bucket] and len(selected) < max_cases:
                candidate = cases_by_bucket[bucket].pop(0)
                if candidate not in selected:
                    selected.append(candidate)
    return selected


def audit_cases_with_llm(
    cases: list[AuditCase],
    checker_configs: dict[str, dict],
    provider: LLMProvider,
    assignment_context: AssignmentContext | str = "",
    max_workers: int = 4,
    progress_callback=None,
) -> list[AuditResult]:
    results = []
    if isinstance(assignment_context, str):
        assignment_context = AssignmentContext(text=assignment_context)
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {
            executor.submit(_audit_one_case, case, checker_configs.get(case.question, {}), provider, assignment_context): case
            for case in cases
        }
        total = len(futures)
        for future in as_completed(futures):
            result = future.result()
            results.append(result)
            if progress_callback:
                progress_callback(result, len(results), total)
    return sorted(results, key=lambda item: (item.question, item.student_id))


def _audit_one_case(case: AuditCase, checker_config: dict, provider: LLMProvider, assignment_context: AssignmentContext) -> AuditResult:
    focused_context = assignment_context_for_question(assignment_context, case.question)
    prompt = build_audit_prompt(case, checker_config, focused_context.text, focused_context.images)
    try:
        primary = _complete_audit_json_with_retry(provider, prompt, focused_context.images)
        primary_decision = _derive_audit_decision(primary, case, focused_context.text)
        responses = [primary]
        decisions = [primary_decision]
        if _needs_challenger_audit(case):
            challenger_prompt = build_audit_prompt(
                case,
                checker_config,
                focused_context.text,
                focused_context.images,
                audit_pass="independent_challenger",
            )
            challenger = _complete_audit_json_with_retry(provider, challenger_prompt, focused_context.images)
            responses.append(challenger)
            decisions.append(_derive_audit_decision(challenger, case, focused_context.text))

        if len(decisions) == 2 and (
            decisions[0]["status"] != decisions[1]["status"]
            or decisions[0]["checker_behavior"] != decisions[1]["checker_behavior"]
        ):
            status = "uncertain"
            verdict = "uncertain"
            risk = "high"
            reason = "Independent audit passes disagreed; checker verification requires human review."
            semantic_assessment = "unclear"
            format_requirement = "unclear"
            checker_behavior = "unclear"
            evidence = " | ".join(str(response.get("evidence", "")) for response in responses if response.get("evidence"))
        else:
            decision = decisions[0]
            response = responses[0]
            status = decision["status"]
            verdict = decision["verdict"]
            risk = str(response.get("risk", "medium"))
            reason = str(response.get("reason", ""))
            semantic_assessment = decision["semantic_assessment"]
            format_requirement = decision["format_requirement"]
            checker_behavior = decision["checker_behavior"]
            evidence = str(response.get("evidence", ""))
    except Exception as exc:
        verdict = "error"
        risk = "high"
        reason = str(exc)
        status = "error"
        semantic_assessment = "unclear"
        format_requirement = "unclear"
        checker_behavior = "unclear"
        evidence = ""
        responses = []
    return AuditResult(
        case.student_id,
        case.question,
        status,
        verdict,
        risk,
        reason,
        semantic_assessment,
        format_requirement,
        checker_behavior,
        evidence,
        max(1, len(responses)),
        tuple(_audit_failed_examples(case.grade_text)),
    )


def _derive_audit_decision(response: dict[str, Any], case: AuditCase, assignment_text: str) -> dict[str, str]:
    semantic = str(response.get("semantic_assessment", "unclear"))
    format_requirement = str(response.get("format_requirement", "unclear"))
    behavior = str(response.get("checker_behavior", "unclear"))
    evidence = str(response.get("evidence", "")).strip()
    legacy_verdict = str(response.get("verdict", "uncertain"))
    extraction_failure = _audit_has_extraction_only_failure(case.grade_text)

    explicit_is_grounded = (
        format_requirement == "explicit"
        and bool(evidence)
        and _evidence_is_grounded(evidence, assignment_text)
    )
    if format_requirement == "explicit" and not explicit_is_grounded:
        format_requirement = "unclear"

    deducted = float(case.score) < 99.999
    if behavior in {"false_reject", "false_accept"}:
        status = "flagged"
        verdict = "flagged"
    elif extraction_failure and semantic == "equivalent" and format_requirement != "explicit":
        status = "flagged"
        verdict = "flagged"
        behavior = "false_reject"
    elif not deducted and semantic == "genuine_error":
        status = "flagged"
        verdict = "flagged"
        behavior = "false_accept"
    elif behavior == "correct" and (
        not deducted or semantic == "genuine_error" or explicit_is_grounded
    ):
        status = "passed"
        verdict = "looks_correct"
    elif (
        "semantic_assessment" not in response
        and legacy_verdict == "looks_correct"
        and not extraction_failure
    ):
        status = "passed"
        verdict = "looks_correct"
    else:
        status = "uncertain"
        verdict = "uncertain"
        if behavior not in {"false_reject", "false_accept", "correct"}:
            behavior = "unclear"

    return {
        "status": status,
        "verdict": verdict,
        "semantic_assessment": semantic,
        "format_requirement": format_requirement,
        "checker_behavior": behavior,
    }


def _needs_challenger_audit(case: AuditCase) -> bool:
    return float(case.score) == 0 or _audit_has_extraction_only_failure(case.grade_text)


def _audit_has_extraction_only_failure(grade_text: str) -> bool:
    lowered = str(grade_text or "").lower()
    extraction_markers = (
        "[missing_anchor]",
        "[missing_label]",
        "anchor '",
        "could not find labeled value",
        "none of the configured boolean aliases",
    )
    return any(marker in lowered for marker in extraction_markers)


def _evidence_is_grounded(evidence: str, assignment_text: str) -> bool:
    normalized_evidence = " ".join(str(evidence).lower().strip(" \"'").split())
    normalized_assignment = " ".join(str(assignment_text).lower().split())
    if len(normalized_evidence) < 8 or not normalized_assignment:
        return False
    if normalized_evidence in normalized_assignment:
        return True
    evidence_words = {word for word in re.findall(r"[a-z0-9]+", normalized_evidence) if len(word) > 3}
    assignment_words = set(re.findall(r"[a-z0-9]+", normalized_assignment))
    return len(evidence_words) >= 3 and len(evidence_words & assignment_words) / len(evidence_words) >= 0.8


def _complete_audit_json_with_retry(
    provider: LLMProvider,
    prompt: str,
    images: list[AssignmentImage] | tuple[AssignmentImage, ...] | None = None,
) -> dict:
    try:
        return complete_json_with_schema(provider, prompt, images, AUDIT_RESPONSE_SCHEMA)
    except Exception as first_error:
        retry_prompt = (
            f"{prompt}\n\n"
            "The previous response could not be parsed as JSON. Return one valid JSON object only, with no markdown, "
            "no comments, and no trailing text. Required keys: verdict, semantic_assessment, format_requirement, "
            "checker_behavior, risk, reason, evidence."
        )
        try:
            return complete_json_with_schema(provider, retry_prompt, images, AUDIT_RESPONSE_SCHEMA)
        except Exception as second_error:
            raise RuntimeError(f"{first_error}; retry failed: {second_error}") from second_error


def build_audit_prompt(
    case: AuditCase,
    checker_config: dict,
    assignment_text: str = "",
    assignment_images: list[AssignmentImage] | tuple[AssignmentImage, ...] | None = None,
    audit_pass: str = "primary",
) -> str:
    grade_evidence = _compact_audit_text(case.grade_text)
    output_evidence = _compact_audit_text(case.output_text)
    reference_evidence = _compact_audit_text(case.reference_output_text)
    return json.dumps(
        {
            "task": "audit_score",
            "audit_rubric_version": AUDIT_RUBRIC_VERSION,
            "audit_pass": audit_pass,
            "role": "You are auditing deterministic C homework grading results. You do not assign a new grade.",
            "student_id": case.student_id,
            "question": case.question,
            "target_question_number": question_number_from_name(case.question),
            "assignment_text_for_selected_question_optional": assignment_text,
            "assignment_images_for_selected_question_optional": [image.label for image in assignment_images or []],
            "eval_cases": CHECKER_AUDIT_EVAL_CASES,
            "checker_config": checker_config,
            "assigned_score": case.score,
            "grade_text": grade_evidence["text"],
            "grade_text_evidence": grade_evidence["metadata"],
            "student_output": output_evidence["text"],
            "student_output_evidence": output_evidence["metadata"],
            "reference_output_by_input": reference_evidence["text"],
            "reference_output_evidence": reference_evidence["metadata"],
            "parsed_discrepancy_signals": _audit_discrepancy_signals(case.grade_text),
            "per_question_excel_fields": case.excel_fields,
            "final_excel_fields": case.final_fields,
            "instructions": (
                f"Focus only on {case.question}. If assignment text or images contain multiple questions, ignore other "
                "question sections. Check whether the assigned score and Excel fields are consistent with the grade text, "
                "student output, checker configuration, and selected-question assignment intent. Verify comments, penalties, compile-error flags, timeout "
                "fields, wrong-input fields, structural recursion/loop check fields, structural penalties, and final weighted grade fields "
                "when present. Do not change the grade. "
                "Audit fairness as well as bookkeeping: a deduction being consistent with the checker's own pass/fail "
                "counts does not make it correct, because the checker itself may be defective. For every deduction, read "
                "the expected-vs-actual discrepancy evidence in the grade text and judge on the merits whether the "
                "student's output genuinely differs in content (wrong numbers, missing required facts, crashes) or only "
                "in semantically equivalent phrasing or formatting. If failures come only from equivalent phrasing or "
                "formatting the assignment does not explicitly require, return flagged and name the checker as the cause. "
                "The application may compact exceptionally long grade or output evidence. When evidence metadata says "
                "application_compacted=true, the omitted middle is an application size limit, not evidence that the student's "
                "program crashed or truncated its output. The source beginning and ending are both retained; use the ending, "
                "runtime/timeout fields, and grade evidence to judge actual completion. "
                "Return looks_correct only when the fields are internally consistent, every output-comparison deduction "
                "reflects a genuine content mistake by the student, and configured policy penalties (compile repair, "
                "structural, submission) follow their documented rules. Return flagged for likely scoring/Excel mistakes "
                "or checker defects that penalize semantically equivalent output. Return uncertain when more human "
                "review is needed. Independently classify semantic_assessment, format_requirement, and checker_behavior. "
                "A missing parser label or anchor is not itself a student semantic error. Claim format_requirement=explicit "
                "only when evidence quotes or closely paraphrases a requirement in the selected assignment context. "
                "Also detect false_accept when a passing/high score hides a genuine semantic error."
            ),
            "response_schema": {
                "verdict": "looks_correct | flagged | uncertain",
                "semantic_assessment": "equivalent | genuine_error | unclear",
                "format_requirement": "explicit | not_explicit | unclear",
                "checker_behavior": "correct | false_reject | false_accept | unclear",
                "risk": "low | medium | high",
                "reason": "one short plain-text sentence under 350 characters; no newlines",
                "evidence": "short assignment quote or concrete expected-vs-actual fact; empty only when evidence is unavailable",
            },
            "strict_json_rules": (
                "Return exactly one JSON object with exactly the keys verdict, semantic_assessment, format_requirement, "
                "checker_behavior, risk, reason, and evidence. "
                "The reason value must be one JSON string: escape quotes and do not include markdown, lists, nested objects, or line breaks."
            ),
        },
        ensure_ascii=False,
        indent=2,
        default=str,
    )


def _audit_discrepancy_signals(grade_text: str) -> dict[str, Any]:
    source = str(grade_text or "")
    reasons = re.findall(r"^Semantic Reason:\s*(.*)$", source, re.MULTILINE)
    expected_canonical = re.findall(r"^Expected Canonical:\s*(.*)$", source, re.MULTILINE)
    actual_canonical = re.findall(r"^Actual Canonical:\s*(.*)$", source, re.MULTILINE)
    return {
        "extraction_only_failure": _audit_has_extraction_only_failure(source),
        "semantic_reasons": reasons[:12],
        "expected_canonical": expected_canonical[:12],
        "actual_canonical": actual_canonical[:12],
    }


def _audit_failed_examples(grade_text: str, max_examples: int = 3) -> list[dict[str, str]]:
    examples = []
    blocks = re.split(r"\n(?=Input:\s*)", str(grade_text or ""))
    for block in blocks:
        input_match = re.search(r"^Input:\s*(.*)$", block, re.MULTILINE)
        expected_match = re.search(r"^Expected:\s*(.*?)(?=^Actual:)", block, re.MULTILINE | re.DOTALL)
        actual_match = re.search(
            r"^Actual:\s*(.*?)(?=^Semantic Reason:|^Expected Canonical:|\Z)",
            block,
            re.MULTILINE | re.DOTALL,
        )
        reason_match = re.search(r"^Semantic Reason:\s*(.*)$", block, re.MULTILINE)
        if not input_match or not expected_match or not actual_match:
            continue
        examples.append(
            {
                "input": input_match.group(1).strip(),
                "expected_output": expected_match.group(1).strip()[:4000],
                "actual_output": actual_match.group(1).strip()[:4000],
                "reason": reason_match.group(1).strip() if reason_match else "",
            }
        )
        if len(examples) >= max_examples:
            break
    return examples


def _compact_audit_text(text: str, max_chars: int = MAX_AUDIT_TEXT_CHARS) -> dict[str, Any]:
    source = str(text or "")
    metadata = {
        "application_compacted": False,
        "source_character_count": len(source),
        "audit_character_limit": max_chars,
        "source_start_included": True,
        "source_end_included": True,
        "note": "Evidence is complete; the application did not truncate it.",
    }
    if len(source) <= max_chars:
        return {"text": source, "metadata": metadata}

    marker_template = (
        "\n\n[AUDIT APPLICATION COMPACTION: {omitted} middle characters omitted because the source exceeded "
        f"{max_chars} characters. This is not student-side truncation. The source ending follows.]\n\n"
    )
    omitted = len(source)
    while True:
        marker = marker_template.format(omitted=omitted)
        available = max(0, max_chars - len(marker))
        start_chars = available // 2
        end_chars = available - start_chars
        updated_omitted = len(source) - start_chars - end_chars
        if updated_omitted == omitted:
            break
        omitted = updated_omitted

    ending = source[-end_chars:] if end_chars else ""
    compacted = source[:start_chars] + marker + ending
    metadata.update(
        {
            "application_compacted": True,
            "omitted_middle_character_count": omitted,
            "note": (
                "The application omitted only the middle due to the audit size limit. "
                "Do not treat this as evidence of student-side truncation or a crash."
            ),
        }
    )
    return {"text": compacted, "metadata": metadata}


def _read_excel_if_exists(path: str) -> pd.DataFrame:
    if not os.path.exists(path):
        return pd.DataFrame()
    return pd.read_excel(path, dtype={"ID_number": str}).fillna("")


def _rows_by_id(df: pd.DataFrame) -> dict[str, dict]:
    if df.empty or "ID_number" not in df.columns:
        return {}
    return {str(row["ID_number"]): row.to_dict() for _, row in df.iterrows()}


def _read_optional_text(path: str) -> str:
    if not os.path.exists(path):
        return ""
    with open(path, "r", encoding="utf-8", errors="ignore") as text_file:
        return text_file.read()


def _case_buckets(case: AuditCase) -> list[str]:
    buckets = []
    lowered_grade = case.grade_text.lower()
    if any(
        marker in lowered_grade
        for marker in ("[missing_anchor]", "[missing_label]", "anchor '", "could not find labeled value")
    ):
        buckets.append("extraction_failure")
    if any(
        marker in lowered_grade
        for marker in ("value mismatch", "status mismatch", "sequence mismatch", "check ")
    ):
        buckets.append("semantic_mismatch")
    if case.score >= 100:
        buckets.append("perfect")
    elif case.score >= 80:
        buckets.append("high")
    elif case.score >= 50:
        buckets.append("medium")
    elif case.score > 0:
        buckets.append("low")
    else:
        buckets.append("zero")
    if case.final_fields.get("Penalty Applied"):
        buckets.append("penalty")
    if case.excel_fields.get("Compilation_Error") or case.excel_fields.get("Timeouts"):
        buckets.append("compile_timeout")
    return buckets
